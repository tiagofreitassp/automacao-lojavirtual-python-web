import os, shutil
from time import sleep

from behave.formatter import null
from docx import Document
from docx.shared import Inches

dir = '../evidencias/'
dirMyStore = '../evidencias/MyStoreOnline'

class myStoreDriverWeb():
    def criarDriver(self):
        self.driver.maximize_window()
        self.driver.get("http://automationpractice.com/index.php?")
        self.driver.implicitly_wait(12)

    def criarPastaEvidencia(self, nPasta):
            d = nPasta
            if os.path.exists(d) == True:
                print("")
                print("Diretório já existe")
                #os.rmdir(d)
                shutil.rmtree(d)
                print('Diretório apagado')
                os.makedirs(d)
                print('Diretório recriado')
            else:
                print("")
                print("Diretório não existe")
                os.makedirs(d)
                print('Diretório criado')

    def gerarScreenshot(self, nPasta, nEvidencia):
        sleep(1)
        self.driver.get_screenshot_as_file(nPasta + "/" + nEvidencia + ".png")
        sleep(1)

    def criarDocumentoDeEvidencia(self, diretorioEvidencia,id ,nomeEvidencia):
        try:
            document = Document()

            h = document.add_heading('MyStore Loja Virtual', 0)
            p = document.add_paragraph("Cenário: "+nomeEvidencia)

            document.add_paragraph(id+'_Tela01')
            document.add_picture(diretorioEvidencia + '/Ev1.png', width=Inches(6.16))
            #document.add_page_break()

            document.add_paragraph(id+'_Tela02')
            document.add_picture(diretorioEvidencia + '/Ev2.png', width=Inches(6.16))
            #document.add_page_break()

            document.add_paragraph(id + '_Tela03')
            document.add_picture(diretorioEvidencia + '/Ev3.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela04')
            document.add_picture(diretorioEvidencia + '/Ev4.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela05')
            document.add_picture(diretorioEvidencia + '/Ev5.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela06')
            document.add_picture(diretorioEvidencia + '/Ev6.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela07')
            document.add_picture(diretorioEvidencia + '/Ev7.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela08')
            document.add_picture(diretorioEvidencia + '/Ev8.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela09')
            document.add_picture(diretorioEvidencia + '/Ev9.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela010')
            document.add_picture(diretorioEvidencia + '/Ev10.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela011')
            document.add_picture(diretorioEvidencia + '/Ev11.png', width=Inches(6.16))

            document.add_paragraph(id + '_Tela012')
            document.add_picture(diretorioEvidencia + '/Ev12.png', width=Inches(6.16))

            document.save(diretorioEvidencia + '/' + nomeEvidencia + '.docx')
            print("Documento com as evidencias gerada com sucesso!")
        except:
            e = Exception
            print(e)
            print("Não foi possivel criar o documento com as evidencias!")

    def fecharDriver(self):
        if self.driver != null:
            self.driver.quit()
        print("Driver finalizado com sucesso!")